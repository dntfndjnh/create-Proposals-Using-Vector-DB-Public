# 설치 필요:
# pip install streamlit transformers sentencepiece faiss-cpu python-docx PyMuPDF scikit-learn keybert

import os
import fitz
import docx
import numpy as np
import faiss
import pickle
import streamlit as st
from sklearn.metrics.pairwise import cosine_similarity
from keybert import KeyBERT
import torch
from transformers import AutoTokenizer, AutoModel

# --- 1️⃣ 문서 읽기 함수 ---
def read_document_paragraphs(file_path_or_file):
    paragraphs = []
    if hasattr(file_path_or_file, "read"):
        # UploadedFile 객체 처리
        if file_path_or_file.name.endswith(".pdf"):
            doc = fitz.open(stream=file_path_or_file.read(), filetype="pdf")
            for page in doc:
                text = page.get_text("text")
                for p in text.split("\n"):
                    if p.strip():
                        paragraphs.append(p.strip())
        elif file_path_or_file.name.endswith(".docx"):
            docx_file = docx.Document(file_path_or_file)
            for p in docx_file.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())
    else:
        # 로컬 경로 처리
        if file_path_or_file.endswith(".pdf"):
            doc = fitz.open(file_path_or_file)
            for page in doc:
                text = page.get_text("text")
                for p in text.split("\n"):
                    if p.strip():
                        paragraphs.append(p.strip())
        elif file_path_or_file.endswith(".docx"):
            if os.path.basename(file_path_or_file).startswith("~$"):
                return []
            d = docx.Document(file_path_or_file)
            for p in d.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())
    return paragraphs

# --- 2️⃣ Streamlit 설정 ---
st.set_page_config(page_title="Document Search & Keyword System", layout="wide")
st.title("문서 검색 및 키워드 추출 시스템. TEAM TechTree")
st.info("문서를 업로드하거나 documents 폴더에 넣으면 자동으로 벡터화됩니다.")

status_message = st.empty()
status_message.info("모델 로드 중... (잠시 기다려주세요)")

# --- 3️⃣ Hugging Face LaBSE 모델 CPU 로드 ---
model_name = "sentence-transformers/LaBSE"
tokenizer = AutoTokenizer.from_pretrained(model_name)
hf_model = AutoModel.from_pretrained(model_name)
hf_model.eval()
hf_model.to("cpu")

# 문장 임베딩 생성 함수
@torch.no_grad()
def get_embedding(text):
    inputs = tokenizer(text, return_tensors="pt", truncation=True, padding=True)
    outputs = hf_model(**inputs)
    # CLS 토큰 벡터 평균
    embeddings = outputs.last_hidden_state.mean(dim=1)
    return embeddings.squeeze().numpy()

# KeyBERT 모델
kw_model = KeyBERT(model=None)  # transformers 모델 직접 사용 예정

stopwords_ko = ["은", "는", "이", "가", "의", "에", "을", "를", "와", "과", "도", "로", "으로"]

status_message.success("모델 로드 완료!")

# --- 4️⃣ FAISS DB 경로 ---
index_file = "vector_index.faiss"
data_file = "doc_data.pkl"

# --- 5️⃣ DB 로드 또는 새로 생성 ---
status_message.info("벡터 DB 준비 중...")
embedding_dim = hf_model.config.hidden_size
if os.path.exists(index_file) and os.path.exists(data_file):
    index = faiss.read_index(index_file)
    with open(data_file, "rb") as f:
        data = pickle.load(f)
        doc_names = data["names"]
        doc_paragraphs = data["paragraphs"]
        doc_embeddings = data["embeddings"]
        doc_keywords = data["keywords"]
    status_message.success("기존 벡터 DB 로드 완료!")
else:
    index = faiss.IndexFlatL2(embedding_dim)
    doc_names, doc_paragraphs, doc_embeddings, doc_keywords = [], [], [], []
    status_message.success("새 벡터 DB 생성 완료!")

# --- 6️⃣ 문서 처리 함수 ---
def process_file(file_path_or_file, file_name, progress_bar=None, progress_offset=0, total_paragraphs=1):
    paragraphs = read_document_paragraphs(file_path_or_file)
    for i, p in enumerate(paragraphs):
        if (file_name, i) in [(d[0], d[1]) for d in doc_paragraphs]:
            continue

        emb = get_embedding(p)
        index.add(np.array([emb], dtype="float32"))
        doc_names.append(file_name)
        doc_paragraphs.append((file_name, i))
        doc_embeddings.append(emb)

        keywords = kw_model.extract_keywords(
            p, keyphrase_ngram_range=(1, 2), stop_words="english", top_n=10
        )
        keywords_list = [kw for kw, score in keywords if not any(sw in kw for sw in stopwords_ko)]
        doc_keywords.append(keywords_list)

        if progress_bar is not None:
            value = min((progress_offset + i + 1) / total_paragraphs, 1.0)
            progress_bar.progress(value)

# --- 7️⃣ documents 폴더 처리 ---
if not os.path.exists("./documents"):
    os.makedirs("./documents")

doc_files = [f for f in os.listdir("./documents") if f.endswith(".pdf") or f.endswith(".docx")]
if doc_files:
    status_message.info("documents 폴더 문서 벡터화 중...")
    progress_bar = st.progress(0)
    total_paragraphs = sum(len(read_document_paragraphs(os.path.join("./documents", f))) for f in doc_files)
    paragraph_offset = 0
    for file_name in doc_files:
        file_path = os.path.join("./documents", file_name)
        process_file(file_path, file_name, progress_bar=progress_bar, progress_offset=paragraph_offset, total_paragraphs=total_paragraphs)
        paragraph_offset += len(read_document_paragraphs(file_path))
    progress_bar.empty()
    status_message.success("documents 폴더 문서 벡터화 완료!")

# --- 8️⃣ Streamlit 업로드 처리 ---
with st.expander("문서 업로드 및 벡터화", expanded=True):
    uploaded_files = st.file_uploader("문서를 선택하세요 (.pdf 또는 .docx)", accept_multiple_files=True)
    if uploaded_files:
        progress_bar = st.progress(0)
        total_paragraphs = 0
        temp_paths = []

        for file in uploaded_files:
            file_name = file.name
            if file_name.startswith("~$") or not (file_name.endswith(".pdf") or file_name.endswith(".docx")):
                continue
            temp_path = os.path.join("./documents", file_name)
            with open(temp_path, "wb") as f:
                f.write(file.getbuffer())
            temp_paths.append((file_name, temp_path))
            total_paragraphs += len(read_document_paragraphs(temp_path))

        paragraph_offset = 0
        for file_name, temp_path in temp_paths:
            process_file(
                file_path_or_file=temp_path,
                file_name=file_name,
                progress_bar=progress_bar,
                progress_offset=paragraph_offset,
                total_paragraphs=total_paragraphs
            )
            paragraph_offset += len(read_document_paragraphs(temp_path))

        progress_bar.empty()
        st.success("업로드된 문서 벡터화 및 키워드 저장 완료!")

# --- 9️⃣ DB 저장 ---
faiss.write_index(index, index_file)
with open(data_file, "wb") as f:
    pickle.dump({
        "names": doc_names,
        "paragraphs": doc_paragraphs,
        "embeddings": doc_embeddings,
        "keywords": doc_keywords
    }, f)

# --- 🔟 검색 기능 ---
with st.expander("문서 검색", expanded=True):
    query = st.text_input("검색어 입력")
    top_k = st.slider("상위 몇 개 결과를 보여드릴까요?", 1, 10, 5)
    if st.button("검색") and query:
        if len(doc_embeddings) == 0:
            st.warning("분석할 문서가 없습니다.")
        else:
            query_emb = get_embedding(query).reshape(1, -1)
            k = min(top_k, len(doc_embeddings))
            distances, indices = index.search(np.array(query_emb, dtype="float32"), k)

            st.subheader(f"검색 결과 (상위 {k}개)")
            for rank, idx in enumerate(indices[0]):
                sim = cosine_similarity([query_emb[0]], [doc_embeddings[idx]])[0][0]
                doc_file, para_idx = doc_paragraphs[idx]
                st.markdown("---")
                st.markdown(f"**{rank+1}. {doc_file} - 문단 {para_idx}**")
                st.markdown(f"- 유사도: {sim:.4f}, 거리: {distances[0][rank]:.4f}")
                st.markdown(f"- 키워드: {', '.join(doc_keywords[idx])}")
