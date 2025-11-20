# ============================================================
# app.py — 슬라이더 재로드 방지/캐시 최적화 + 로딩 프로그레스바
# ============================================================

import os, pickle, hashlib, faiss, numpy as np, streamlit as st
import fitz, docx
from docx import Document
from sklearn.metrics.pairwise import cosine_similarity
from keybert import KeyBERT
import torch
from transformers import AutoTokenizer, AutoModel
from io import BytesIO
from openai import OpenAI
import streamlit.components.v1 as components
from typing import List, Tuple

# -----------------------
# 페이지 설정
# -----------------------
st.set_page_config(page_title="Document Search & Auto Plan", layout="wide")
st.title("문서 검색 · 키워드 · AI 기획서 생성 — TEAM TechTree")
st.info("문서를 업로드하거나 ./documents 폴더에 넣으면 벡터화됩니다.")

# -----------------------
# OpenAI KEY
# -----------------------
openai_key = st.secrets.get("OPENAI_API_KEY", "") or os.environ.get("OPENAI_API_KEY", "")
client = OpenAI(api_key=openai_key) if openai_key else None
if not client:
    st.warning("⚠️ OpenAI API 키가 설정되지 않았습니다.")

# -----------------------
# 파일 상수
# -----------------------
INDEX_FILE, DATA_FILE, HASH_FILE = "vector_index.faiss", "doc_data.pkl", "doc_hash.pkl"
os.makedirs("./documents", exist_ok=True)

# -----------------------
# 캐시 로드
# -----------------------
@st.cache_resource
def load_hf_model(model_name="sentence-transformers/LaBSE"):
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    hf_model = AutoModel.from_pretrained(model_name)
    hf_model.eval().to("cpu")
    return tokenizer, hf_model

@st.cache_resource
def load_kw_model():
    torch.set_default_device("cpu")
    kw = KeyBERT(model="all-mpnet-base-v2")
    kw.model.embedding_model.to("cpu")
    return kw

@st.cache_resource
def load_faiss_and_db(index_file=INDEX_FILE, data_file=DATA_FILE, embedding_dim=768):
    if os.path.exists(index_file) and os.path.exists(data_file):
        try:
            idx = faiss.read_index(index_file)
            db = pickle.load(open(data_file, "rb"))
            return idx, db.get("names", []), db.get("paragraphs", []), db.get("embeddings", []), db.get("keywords", [])
        except: pass
    return faiss.IndexFlatL2(embedding_dim), [], [], [], []

@st.cache_data
def read_document_paragraphs_cached(path: str) -> List[str]:
    paragraphs = []
    try:
        if path.endswith(".pdf"):
            pdf = fitz.open(path)
            for page in pdf:
                paragraphs += [p.strip() for p in page.get_text("text").split("\n") if p.strip()]
        elif path.endswith(".docx"):
            if not os.path.basename(path).startswith("~$"):
                d = docx.Document(path)
                paragraphs += [p.text.strip() for p in d.paragraphs if p.text.strip()]
    except Exception as e:
        st.error(f"문서 읽기 실패 ({path}): {e}")
    return paragraphs

# -----------------------
# 초기 리소스 로드 + 프로그레스 표시
# -----------------------
status = st.empty()
pb = st.progress(0)
step = 0

status.info("모델/리소스 로드 중...")

step += 1; pb.progress(step/5)
tokenizer, hf_model = load_hf_model()

step += 1; pb.progress(step/5)
kw_model = load_kw_model()

step += 1; pb.progress(step/5)
EMBED_DIM = hf_model.config.hidden_size
index, doc_names, doc_paragraphs, doc_embeddings, doc_keywords = load_faiss_and_db(embedding_dim=EMBED_DIM)

step += 1; pb.progress(step/5)
doc_hash = pickle.load(open(HASH_FILE, "rb")) if os.path.exists(HASH_FILE) else {}

step += 1; pb.progress(step/5)
status.success("리소스 로드 완료")
pb.empty()  # 프로그레스바 제거

# -----------------------
# 유틸
# -----------------------
@torch.no_grad()
def get_embedding(text, tokenizer=tokenizer, model=hf_model):
    if not text: return np.zeros(model.config.hidden_size, dtype="float32")
    inputs = tokenizer(text, return_tensors="pt", truncation=True, padding=True, max_length=512)
    outputs = model(**inputs)
    return outputs.last_hidden_state.mean(dim=1).squeeze().cpu().numpy().astype("float32")

def compute_hash(file_path):
    h = hashlib.sha256()
    with open(file_path, "rb") as f: h.update(f.read())
    return h.hexdigest()

def rebuild_faiss(remove_list: List[int]):
    global index, doc_names, doc_paragraphs, doc_embeddings, doc_keywords
    keep = [i for i in range(len(doc_embeddings)) if i not in remove_list]
    doc_names = [doc_names[i] for i in keep]
    doc_paragraphs = [doc_paragraphs[i] for i in keep]
    doc_embeddings = [doc_embeddings[i] for i in keep]
    doc_keywords = [doc_keywords[i] for i in keep]
    index = faiss.IndexFlatL2(EMBED_DIM)
    if doc_embeddings: index.add(np.array(doc_embeddings, dtype="float32"))

# -----------------------
# 파일 처리 (슬라이더와 독립)
# -----------------------
def process_file(path: str, filename: str):
    global doc_names, doc_paragraphs, doc_embeddings, doc_keywords, doc_hash, index
    paragraphs = read_document_paragraphs_cached(path)
    new_hash = compute_hash(path)
    if filename in doc_hash and doc_hash[filename] == new_hash: return len(paragraphs)
    if filename in doc_names:
        remove_idxs = [i for i, (fn, _) in enumerate(doc_paragraphs) if fn == filename]
        if remove_idxs: rebuild_faiss(remove_idxs)
    for i, p in enumerate(paragraphs):
        emb = get_embedding(p)
        index.add(np.array([emb], dtype="float32"))
        doc_names.append(filename)
        doc_paragraphs.append((filename, i))
        doc_embeddings.append(emb)
        try:
            kws = kw_model.extract_keywords(p, keyphrase_ngram_range=(1, 2), stop_words=None, top_n=8)
            doc_keywords.append([k for k, _ in kws])
        except: doc_keywords.append(p.split()[:8])
    doc_hash[filename] = new_hash
    return len(paragraphs)

def save_db():
    faiss.write_index(index, INDEX_FILE)
    pickle.dump({"names": doc_names, "paragraphs": doc_paragraphs, "embeddings": doc_embeddings, "keywords": doc_keywords}, open(DATA_FILE, "wb"))
    pickle.dump(doc_hash, open(HASH_FILE, "wb"))
    st.success("DB 저장 완료!")

# -----------------------
# 문서 스캔 (앱 시작 시만) + 프로그레스
# -----------------------
if "docs_scanned" not in st.session_state:
    files = [f for f in os.listdir("./documents") if f.endswith((".pdf", ".docx"))]
    added_count = 0
    pb = st.progress(0)
    for i, f in enumerate(files):
        added_count += process_file(os.path.join("./documents", f), f)
        pb.progress((i+1)/len(files))
    pb.empty()
    st.session_state.docs_scanned = True
    if added_count: st.success(f"문서 벡터화 완료 — {added_count} 문단 추가됨")

# -----------------------
# 업로드 UI
# -----------------------
with st.expander("문서 업로드 및 벡터화", expanded=True):
    uploaded = st.file_uploader("문서 선택 (.pdf/.docx)", accept_multiple_files=True)
    if uploaded:
        for uf in uploaded:
            target = os.path.join("./documents", uf.name)
            with open(target, "wb") as f: f.write(uf.getbuffer())
            process_file(target, uf.name)
        st.success("업로드 문서 벡터화 완료")
if st.button("DB 저장"): save_db()

# -----------------------
# 검색 UI
# -----------------------
if "kw_selection" not in st.session_state: st.session_state.kw_selection = []
if "last_search" not in st.session_state: st.session_state.last_search = []

with st.expander("문서 검색", expanded=True):
    query = st.text_input("검색어 입력")
    top_k = st.slider("상위 결과 수", 1, 10, 5)
    if st.button("검색 실행"):
        if not doc_embeddings: st.warning("DB가 비어 있습니다.")
        else:
            q_emb = get_embedding(query).reshape(1, -1)
            k = min(top_k, len(doc_embeddings))
            dist, idxs = index.search(q_emb.astype("float32"), k)
            results, all_kw = [], []
            for rank, ix in enumerate(idxs[0]):
                sim = cosine_similarity([q_emb[0]], [doc_embeddings[ix]])[0][0]
                fn, pidx = doc_paragraphs[ix]; kws = doc_keywords[ix]
                all_kw += kws
                results.append({"file": fn, "paragraph_idx": pidx, "similarity": sim, "keywords": kws})
                st.markdown(f"**{rank+1}. {fn} (문단 {pidx})**")
                st.write(f"유사도: {sim:.4f} | 키워드: {', '.join(kws)}")
            st.session_state.last_search = results
            st.session_state.kw_selection = list(dict.fromkeys(all_kw))[:6]

    if st.session_state.last_search:
        options = list(dict.fromkeys([kw for r in st.session_state.last_search for kw in r["keywords"]]))
        selected_keywords = st.multiselect("기획서에 넣을 키워드", options=options, default=st.session_state.kw_selection)
        if selected_keywords != st.session_state.kw_selection:
            st.session_state.kw_selection = selected_keywords

# -----------------------
# 기획서 생성
# -----------------------
def generate_project_plan_openai(keywords: List[str], notes: str = "") -> Tuple[str, str]:
    if not client: raise RuntimeError("OpenAI client not configured")
    kw = ", ".join(keywords)
    prompt = f"""
당신은 소프트웨어 기획 전문가입니다.
아래 키워드를 기반으로 AI 프로젝트 기획서를 작성하세요.

키워드: {kw}
추가요청: {notes}

[PLAN]
- 프로젝트 개요
- 문제 정의 & 목표
- 핵심 기능(우선순위)
- 아키텍처
- 개발 로드맵
- 기대 효과

[MERMAID] flowchart 코드 포함
"""
    res = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.2, max_tokens=1200)
    text = res.choices[0].message.content
    if "[MERMAID]" in text: plan, mermaid = text.split("[MERMAID]"); return plan.replace("[PLAN]", "").strip(), mermaid.strip()
    return text, None

with st.expander("AI 자동 기획서 생성", expanded=True):
    kw_input = st.text_area("사용할 키워드", value=", ".join(st.session_state.kw_selection))
    notes = st.text_area("추가 요청")
    if st.button("기획서 생성"):
        if not openai_key: st.error("API 키가 없습니다.")
        elif not kw_input.strip(): st.warning("키워드를 입력하세요.")
        else:
            with st.spinner("생성 중..."):
                kws = [k.strip() for k in kw_input.split(",") if k.strip()]
                plan, mermaid = generate_project_plan_openai(kws, notes)
                st.markdown(plan)
                if mermaid:
                    components.html(f"""
<div class="mermaid">{mermaid}</div>
<script type="module">
import mermaid from "https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs";
mermaid.initialize({{ startOnLoad: false }});
document.querySelectorAll('.mermaid').forEach(el => mermaid.init(undefined, el));
</script>
""", height=500)
                def create_docx(plan_text, mermaid_code=None):
                    doc = Document(); doc.add_heading("AI Project Plan", 0)
                    for line in plan_text.split("\n"): doc.add_paragraph(line)
                    if mermaid_code: 
                        doc.add_heading("Mermaid Code", 1)
                        for line in mermaid_code.split("\n"): doc.add_paragraph(line)
                    bio = BytesIO(); doc.save(bio); bio.seek(0); return bio
                st.download_button("DOCX 다운로드", data=create_docx(plan, mermaid), file_name="project_plan.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                if mermaid:
                    st.download_button("Mermaid(.mmd) 다운로드", data=mermaid, file_name="diagram.mmd", mime="text/plain")

st.caption("DB 저장을 눌러 벡터 DB를 보존하세요.")
